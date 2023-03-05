from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromiumService
from webdriver_manager.chrome import ChromeDriverManager
from webdriver_manager.core.utils import ChromeType
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.service import Service as ChromeService
from PyPDF2 import PdfReader
import chromedriver_binary 
from datetime import datetime
import openai
import os
from urllib.request import Request, urlopen
from bs4 import BeautifulSoup
import requests
import smtplib
import ssl
from email.message import EmailMessage
import datetime
import re
import smtplib
import dns.resolver
import socket
from dotenv import load_dotenv
from docx import Document

load_dotenv()
openai.api_key = os.getenv('OPENAI_API_KEY')


class TomideBeautifulSoupUtils:

    def __init__(self, url, type, scroll):
        self.url = url
        self.type = type
        self.scroll = scroll

    @staticmethod
    def get_classes(soup):
        class_list = []
        tags = {tag for tag in soup.find_all()}
        for tag in tags:
            if tag.has_attr("class"):
                if len(tag['class']) != 0:
                    if tag['class'][0] not in class_list:
                        class_list.append(tag['class'][0])
        return class_list

    @classmethod
    def get_all_links(cls, soup, url):
        all_links = soup.find_all('a')
        all_links = [link['href'] for link in all_links]

        for link in all_links:
            if link[0] == '/':
                all_links.append(url + link.replace('/', ''))

        # remove links starting with / or #
        all_links = [
            link for link in all_links if link[0] != '/' and link[0] != '#'
        ]

        internal_links = [
            internal_link for internal_link in all_links
            if url in internal_link
        ]
        external_links = [
            external_link for external_link in all_links
            if url not in external_link
        ]
        socials = [
            'facebook', 'twitter', 'instagram', 'linkedin', 'youtube',
            'pinterest', 'tumblr', 'reddit', 'snapchat', 'whatsapp',
            'telegram', 't.me', 'wa.me'
        ]
        social_media_links = []
        for link in all_links:
            for social in socials:
                if social in link:
                    social_media_links.append(link)
        return {
            "internal_links": internal_links,
            "external_links": external_links,
            "social_media_links": social_media_links
        }

    @classmethod
    def tomide_bs4_make_soup(cls, url, type):
        soup = None

        if type == "static":
            req = Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            webpage = urlopen(req).read()
            soup = BeautifulSoup(webpage, 'html.parser')
        elif type == "chromium":
            driver = webdriver.Chrome(service=ChromiumService(
                ChromeDriverManager(
                    chrome_type=ChromeType.CHROMIUM).install()))
            driver.get(url)
            soup = BeautifulSoup(driver.page_source, 'html.parser')
            return soup, cls.get_all_links(soup, url)
        elif type == "firefox":
            from selenium.webdriver.firefox.firefox_binary import FirefoxBinary
            firefox_binary = FirefoxBinary()
            driver = webdriver.Firefox(firefox_binary=firefox_binary)
            soup = BeautifulSoup(driver.page_source, 'html.parser')
            driver.quit()
            return soup, cls.get_all_links(soup, url)

        else:
            options = webdriver.ChromeOptions()
            options.headless = True
            options.page_load_strategy = 'none'
            service = ChromeService('./chromedriver')

            driver = webdriver.Chrome(
                service=service,
                options=options,
                executable_path='./chromedriver')
            # if scroll == True:
            #     SCROLL_PAUSE_TIME = 10
            #     last_height = driver.execute_script(
            #         "return document.body.scrollHeight")  # Get scroll height
            #     while True:
            #         driver.execute_script(  # Scroll down to bottom
            #             "window.scrollTo(0, document.body.scrollHeight);")
            #         time.sleep(SCROLL_PAUSE_TIME)  # Wait to load page
            #         new_height = driver.execute_script(  # Calculate new scroll height and compare with last scroll height
            #             "return document.body.scrollHeight")
            #         if new_height == last_height:
            #             break
            #         last_height = new_height

            driver.get(url)
            soup = BeautifulSoup(driver.page_source, 'html.parser')
            driver.quit()

        return soup, cls.get_all_links(soup, url)


class TomsEmailUtilities:

    def __init__(self, firstname, lastname, company_domain):
        self.firstname = firstname
        self.lastname = lastname
        self.company_domain = company_domain

    @staticmethod
    def email_extractor(line):
        return re.findall(r'[\w.+-]+@[\w-]+\.[\w.-]+', line)

    @staticmethod
    def mail_checker(email):
        domain_name = email.split('@')[1]
        records = dns.resolver.query(domain_name, 'MX')
        mxRecord = records[0].exchange
        mxRecord = str(mxRecord)
        host = socket.gethostname()
        server = smtplib.SMTP()
        server.set_debuglevel(0)
        server.connect(mxRecord)
        server.helo(host)
        server.mail('me@domain.com')
        code, message = server.rcpt(str(email))
        server.quit()
        if code == 250:
            print(email, "is valid")
            return email

    @classmethod
    def emailGenerator(cls, firstname, lastname, company_domain):
        if '/' in company_domain:
            company_domain = company_domain.replace('/', '')
        if 'https:' in company_domain:
            company_domain = company_domain.replace('https:', '')
        if 'http:' in company_domain:
            company_domain = company_domain.replace('http:', '')
        if 'www.' in company_domain:
            company_domain = company_domain.replace('www.', '')

        print("Generating email list for", firstname, lastname)
        emailList = []

        emailList.append(firstname + "@" + company_domain)
        emailList.append(lastname + "@" + company_domain)

        emailList.append(firstname + lastname + "@" + company_domain)
        emailList.append(lastname + firstname + "@" + company_domain)

        emailList.append(firstname + "." + lastname + "@" +
                         company_domain)  # tomide.adeoye@merislabs.com
        emailList.append(lastname + "." + firstname + "@" +
                         company_domain)  # adeoye.tomide@merislabs.com

        # tomidea@merislabs.com
        emailList.append(firstname + lastname[0] + "@" + company_domain)
        # atomide@merislabs.com
        emailList.append(lastname[0] + firstname + "@" + company_domain)

        # tadeoye@merislabs.com
        emailList.append(firstname[0] + lastname + "@" + company_domain)
        # adeoyet@merislabs.com
        emailList.append(lastname + firstname[0] + "@" + company_domain)

        # tomidea@merislabs.com
        emailList.append(firstname + "." + lastname[0] + "@" + company_domain)
        # atomide@merislabs.com
        emailList.append(lastname[0] + "." + firstname + "@" + company_domain)

        # tadeoye@merislabs.com
        emailList.append(firstname[0] + "." + lastname + "@" + company_domain)
        # adeoyet@merislabs.com
        emailList.append(lastname + "." + firstname[0] + "@" + company_domain)

        emailList.append(firstname[0] + lastname[0] + "@" +
                         company_domain)  # ta@merislabs.com
        emailList.append(lastname[0] + firstname[0] + "@" +
                         company_domain)  # ta@merislabs.com

        emailList.append(lastname[0] + "." + firstname[0] + "@" +
                         company_domain)  # t.a@meris.com
        # t.a@meris.com
        emailList.append(firstname[0] + "." + lastname[0] + "@" +
                         company_domain)

        valid = []
        for email in emailList:
            # print position of email in array
            print('Validating ',
                  emailList.index(email) + 1, ' of ', len(emailList))
            valid.append(cls.mail_checker(email))

        valid = [item for item in valid if item is not None]

        return valid

    @staticmethod
    def send_email(email_receiver, subject, body, files):
        now = datetime.datetime.now()
        email_sender = os.getenv('email_sender')
        email_password = os.getenv('email_password')
        em = EmailMessage()
        em['From'], em['To'], em[
            'Subject'] = email_sender, email_receiver, subject
        em.set_content(body)

        for path in files:
            with open(path, 'rb') as file:
                file_data = file.read()
                file_name = file.name.split('/')[-1]

            em.add_attachment(file_data,
                              maintype='application',
                              subtype='octet-stream',
                              filename=file_name)

        context = ssl.create_default_context()
        with smtplib.SMTP_SSL('smtp.gmail.com', 465, context=context) as smtp:
            smtp.login(email_sender, email_password)
            smtp.send_message(em)
            smtp.quit()


def google_search(query_for_google):
    search_keys = os.getenv("search_keys").split(
        ",")  # an array of search keys
    searchEngineId = os.getenv("search_engine_id")
    options = {'method': 'get', 'contentType': 'application/json'}

    for search_key in search_keys:
        google_response = requests.get(
            "https://www.googleapis.com/customsearch/v1?key=" + search_key +
            "&q=" + query_for_google + "&cx=" + searchEngineId,
            options).json()
        if 'items' in google_response:
            return google_response['items']
        else:
            if google_response['error']['code'] == 429:
                continue
    print("NO KEYS WORKED", google_response['error']['message'])


def create_document(title, body):

    created_at = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    document = Document()
    document.add_heading(title, 1)
    document.add_paragraph(body)
    document.add_page_break()
    document.add_paragraph('Created at ' + created_at + ' by Tomide Adeoye')

    document_name = f"""{title}-{created_at}.docx"""
    document.save(document_name)

    return document_name


def replace_irrelevant_words(text):
    things_to_replace = [
        'all rights reserved',
        '©',
        '®',
        '™',
        'incognitp',
        '<b>',
        '</b>',
    ]

    for thing in things_to_replace:
        text = text.replace(thing, '')

    return text